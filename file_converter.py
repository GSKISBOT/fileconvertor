"""
File conversion utilities for converting various file formats to DOCX
"""

import logging
import os
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Inches
import PyPDF2
import pytesseract
from PIL import Image
import io
import zipfile
import xml.etree.ElementTree as ET
from utils import cleanup_file, get_file_extension

logger = logging.getLogger(__name__)

class FileConverter:
    def __init__(self):
        self.supported_formats = ['.txt', '.pdf', '.doc', '.docx', '.rtf', '.odt', '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']
    
    def convert_to_docx(self, file_path: str, original_filename: str) -> str:
        """Convert any supported file to DOCX format"""
        try:
            file_extension = get_file_extension(original_filename).lower()
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Extract text based on file type
            if file_extension == '.txt':
                text_content = self._extract_from_txt(file_path)
            elif file_extension == '.pdf':
                text_content = self._extract_from_pdf(file_path)
            elif file_extension in ['.doc', '.docx']:
                text_content = self._extract_from_word(file_path)
            elif file_extension == '.rtf':
                text_content = self._extract_from_rtf(file_path)
            elif file_extension == '.odt':
                text_content = self._extract_from_odt(file_path)
            elif file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
                text_content = self._extract_from_image(file_path)
            else:
                raise ValueError(f"Conversion not implemented for {file_extension}")
            
            # Create DOCX file
            output_path = self._create_docx(text_content, original_filename)
            return output_path
            
        except Exception as e:
            logger.error(f"Error converting file {original_filename}: {e}")
            raise
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text_content = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"Error extracting from PDF: {e}")
            raise ValueError("Failed to extract text from PDF")
        
        if not text_content.strip():
            raise ValueError("No text found in PDF")
        
        return text_content
    
    def _extract_from_word(self, file_path: str) -> str:
        """Extract text from Word document"""
        try:
            doc = Document(file_path)
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            return text_content
        except Exception as e:
            logger.error(f"Error extracting from Word document: {e}")
            raise ValueError("Failed to extract text from Word document")
    
    def _extract_from_rtf(self, file_path: str) -> str:
        """Extract text from RTF file"""
        try:
            # Simple RTF text extraction (basic implementation)
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Remove RTF control sequences (basic cleanup)
            import re
            # Remove RTF control words
            content = re.sub(r'\\[a-z]+\d*', '', content)
            # Remove braces
            content = re.sub(r'[{}]', '', content)
            # Clean up extra whitespace
            content = re.sub(r'\s+', ' ', content)
            
            return content.strip()
        except Exception as e:
            logger.error(f"Error extracting from RTF: {e}")
            raise ValueError("Failed to extract text from RTF file")
    
    def _extract_from_odt(self, file_path: str) -> str:
        """Extract text from ODT file"""
        try:
            text_content = ""
            with zipfile.ZipFile(file_path, 'r') as odt_file:
                content_xml = odt_file.read('content.xml')
                root = ET.fromstring(content_xml)
                
                # Extract text from all text nodes
                for elem in root.iter():
                    if elem.text:
                        text_content += elem.text + " "
                    if elem.tail:
                        text_content += elem.tail + " "
            
            return text_content.strip()
        except Exception as e:
            logger.error(f"Error extracting from ODT: {e}")
            raise ValueError("Failed to extract text from ODT file")
    
    def _extract_from_image(self, file_path: str) -> str:
        """Extract text from image using OCR"""
        try:
            # Open image and perform OCR
            image = Image.open(file_path)
            text_content = pytesseract.image_to_string(image)
            
            if not text_content.strip():
                raise ValueError("No text found in image")
            
            return text_content
        except Exception as e:
            logger.error(f"Error extracting text from image: {e}")
            raise ValueError("Failed to extract text from image. Make sure the image contains readable text.")
    
    def _create_docx(self, text_content: str, original_filename: str) -> str:
        """Create a DOCX file from text content"""
        try:
            # Create new document
            doc = Document()
            
            # Add title
            title = doc.add_heading(f'Converted from {original_filename}', 0)
            
            # Add content
            paragraphs = text_content.split('\n')
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text.strip())
            
            # Save to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                doc.save(temp_file.name)
                return temp_file.name
                
        except Exception as e:
            logger.error(f"Error creating DOCX: {e}")
            raise ValueError("Failed to create DOCX file")
