"""
Translation services for converting text content between languages
"""

import logging
import tempfile
from docx import Document
from googletrans import Translator as GoogleTranslator
from file_converter import FileConverter
from utils import cleanup_file, get_file_extension
from config import SUPPORTED_LANGUAGES

logger = logging.getLogger(__name__)

class Translator:
    def __init__(self):
        self.google_translator = GoogleTranslator()
        self.file_converter = FileConverter()
        self.max_chunk_size = 5000  # Google Translate has character limits
    
    def translate_file(self, file_path: str, target_lang: str, original_filename: str) -> str:
        """Translate file content to target language and return as DOCX"""
        try:
            # First extract text from the file
            text_content = self._extract_text_from_file(file_path, original_filename)
            
            if not text_content.strip():
                raise ValueError("No text content found in file")
            
            # Detect source language
            detection = self.google_translator.detect(text_content[:1000])  # Use first 1000 chars for detection
            source_lang = detection.lang
            confidence = detection.confidence
            
            logger.info(f"Detected source language: {source_lang} (confidence: {confidence})")
            
            # Translate text
            translated_text = self._translate_text(text_content, source_lang, target_lang)
            
            # Create DOCX with translated content
            output_path = self._create_translated_docx(
                translated_text, 
                original_filename, 
                source_lang, 
                target_lang
            )
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error translating file {original_filename}: {e}")
            raise
    
    def _extract_text_from_file(self, file_path: str, original_filename: str) -> str:
        """Extract text from file using FileConverter"""
        file_extension = get_file_extension(original_filename).lower()
        
        if file_extension == '.txt':
            return self.file_converter._extract_from_txt(file_path)
        elif file_extension == '.pdf':
            return self.file_converter._extract_from_pdf(file_path)
        elif file_extension in ['.doc', '.docx']:
            return self.file_converter._extract_from_word(file_path)
        elif file_extension == '.rtf':
            return self.file_converter._extract_from_rtf(file_path)
        elif file_extension == '.odt':
            return self.file_converter._extract_from_odt(file_path)
        elif file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
            return self.file_converter._extract_from_image(file_path)
        else:
            raise ValueError(f"Unsupported file format for translation: {file_extension}")
    
    def _translate_text(self, text: str, source_lang: str, target_lang: str) -> str:
        """Translate text content, handling large texts by chunking"""
        try:
            if len(text) <= self.max_chunk_size:
                # Text is small enough to translate in one go
                result = self.google_translator.translate(
                    text, 
                    src=source_lang, 
                    dest=target_lang
                )
                return result.text
            else:
                # Split text into chunks
                chunks = self._split_text_into_chunks(text)
                translated_chunks = []
                
                for chunk in chunks:
                    if chunk.strip():
                        result = self.google_translator.translate(
                            chunk, 
                            src=source_lang, 
                            dest=target_lang
                        )
                        translated_chunks.append(result.text)
                
                return '\n'.join(translated_chunks)
                
        except Exception as e:
            logger.error(f"Error translating text: {e}")
            raise ValueError("Translation failed. Please try again.")
    
    def _split_text_into_chunks(self, text: str) -> list:
        """Split text into chunks that respect sentence boundaries"""
        chunks = []
        current_chunk = ""
        
        # Split by sentences (simple approach)
        sentences = text.replace('.', '.\n').replace('!', '!\n').replace('?', '?\n').split('\n')
        
        for sentence in sentences:
            if len(current_chunk + sentence) <= self.max_chunk_size:
                current_chunk += sentence + " "
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence + " "
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def _create_translated_docx(self, translated_text: str, original_filename: str, source_lang: str, target_lang: str) -> str:
        """Create a DOCX file with translated content"""
        try:
            # Create new document
            doc = Document()
            
            # Add title with translation info
            source_lang_name = SUPPORTED_LANGUAGES.get(source_lang, source_lang.upper())
            target_lang_name = SUPPORTED_LANGUAGES.get(target_lang, target_lang.upper())
            
            title = doc.add_heading(f'Translated: {original_filename}', 0)
            subtitle = doc.add_paragraph(f'From {source_lang_name} to {target_lang_name}')
            subtitle.style = 'Subtitle'
            
            # Add horizontal line
            doc.add_paragraph('─' * 50)
            
            # Add translated content
            paragraphs = translated_text.split('\n')
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text.strip())
            
            # Save to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                doc.save(temp_file.name)
                return temp_file.name
                
        except Exception as e:
            logger.error(f"Error creating translated DOCX: {e}")
            raise ValueError("Failed to create translated document")
